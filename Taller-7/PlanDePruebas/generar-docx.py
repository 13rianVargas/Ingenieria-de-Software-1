#!/usr/bin/env python3
"""
Generador del Plan de Pruebas — Taller-7
E-Commerce Comercial Konrad

Pipeline: contenido/*.md + 00-portada.yaml → plan-de-pruebas.docx
Requiere: pip install python-docx pyyaml
Ejecutar:  python generar-docx.py
"""

import re
import yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import estilos as E

# ── Rutas ─────────────────────────────────────────────────────────────────────
BASE = Path(__file__).parent
CONTENIDO = BASE / "contenido"
OUTPUT = BASE / "plan-de-pruebas.docx"

SECCIONES = [
    "01-resumen-ejecutivo.md",
    "02-alcance.md",
    "03-estrategia.md",
    "04-criterios.md",
    "05-recursos.md",
    "06-planificacion.md",
    "07-riesgos.md",
    "08-anexos.md",
]


# ── Portada ───────────────────────────────────────────────────────────────────
def build_portada(doc, meta):
    sec = doc.sections[0]
    sec.page_width  = Cm(21.59)   # carta
    sec.page_height = Cm(27.94)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

    # Institucion
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.get("institucion", "").upper())
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = E.COLOR_PRIMARY
    E._safe_font(run)
    E._add_bottom_border(p)
    p.paragraph_format.space_after = Pt(48)

    # Titulo principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.get("titulo", "Plan de Pruebas de Software"))
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = E.COLOR_PRIMARY
    E._safe_font(run)
    p.paragraph_format.space_after = Pt(4)

    # Linea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    E._add_bottom_border(p, size=16)
    p.paragraph_format.space_after = Pt(8)

    # Subtitulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.get("subtitulo", ""))
    run.font.size = Pt(14)
    run.font.color.rgb = E.COLOR_SECONDARY
    E._safe_font(run)
    p.paragraph_format.space_after = Pt(48)

    # Autores
    for autor in meta.get("autores", []):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(autor)
        run.font.size = Pt(11)
        run.font.color.rgb = E.COLOR_TEXT
        E._safe_font(run)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(32)

    # Datos del curso
    datos = [
        meta.get("materia", ""),
        f"Grupo {meta.get('grupo', '')}",
        f"Docente: {meta.get('docente', '')}",
        meta.get("fecha", ""),
    ]
    for dato in datos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dato)
        run.font.size = Pt(10)
        run.font.color.rgb = E.COLOR_MUTED
        E._safe_font(run)
        p.paragraph_format.space_after = Pt(3)

    # Salto de pagina al final de portada
    doc.add_page_break()


# ── Tabla de Contenido ────────────────────────────────────────────────────────
def build_toc(doc):
    E.add_heading(doc, "Tabla de Contenido", level=1)
    E.add_toc_field(doc)
    p = doc.add_paragraph()
    run = p.add_run("(Actualizar con F9 en Word o al abrir el archivo)")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = E.COLOR_MUTED
    doc.add_page_break()


# ── Historial de versiones ────────────────────────────────────────────────────
def build_historial(doc, meta):
    E.add_heading(doc, "Historial de Versiones", level=1)
    E.add_table(
        doc,
        headers=["Fecha", "Version", "Autor", "Descripcion"],
        rows=[
            [meta.get("fecha", ""), meta.get("version", "1.0"),
             " · ".join(a.split()[0] for a in meta.get("autores", [])), "Version inicial del plan de pruebas"],
        ],
        col_widths_cm=[3, 2, 5, 8],
    )
    doc.add_paragraph()


# ── Informacion del Proyecto ──────────────────────────────────────────────────
def build_info_proyecto(doc, meta):
    E.add_heading(doc, "Informacion del Proyecto", level=1)
    filas = [
        ["Empresa / Organizacion", meta.get("institucion", "")],
        ["Proyecto",               meta.get("subtitulo",   "")],
        ["Cliente",                meta.get("cliente",     "Comercial Konrad")],
        ["Fecha de preparacion",   meta.get("fecha",       "")],
        ["Materia",                meta.get("materia",     "")],
        ["Docente",                meta.get("docente",     "")],
        ["Grupo",                  meta.get("grupo",       "")],
        ["Lider de Pruebas",       "Por asignar internamente"],
    ]
    E.add_table(doc, headers=["Campo", "Valor"], rows=filas,
                col_widths_cm=[6, 11])
    doc.add_paragraph()


# ── Parser Markdown basico ────────────────────────────────────────────────────
def parse_markdown_section(doc, md_text):
    """
    Parser simplificado: soporta # headings, tablas pipe, listas - y texto.
    No soporta codigo inline ni negrita en cuerpo (mejora futura).
    """
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading — limpiar backticks, bold, italic y links del texto
        def _clean(t):
            t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", t)  # [text](url) → text
            t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)        # **bold** → text
            t = re.sub(r"\*([^*]+)\*", r"\1", t)             # *italic* → text
            t = re.sub(r"`([^`]+)`", r"\1", t)
            return t.replace("§", "sección").strip()

        if line.startswith("### "):
            E.add_heading(doc, _clean(line[4:]), level=3)
        elif line.startswith("## "):
            E.add_heading(doc, _clean(line[3:]), level=2)
        elif line.startswith("# "):
            E.add_heading(doc, _clean(line[2:]), level=1)

        # Tabla pipe
        elif line.strip().startswith("|") and "|" in line:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not re.match(r"^\s*\|[-| :]+\|\s*$", lines[i]):
                    table_lines.append(lines[i])
                i += 1
            if table_lines:
                _render_table(doc, table_lines)
            continue

        # Lista con -
        elif line.strip().startswith("- "):
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
            _add_inline_runs(p, line.strip()[2:], size_pt=10.5)

        # Lista numerada
        elif re.match(r"^\d+\. ", line.strip()):
            text = re.sub(r"^\d+\. ", "", line.strip())
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
            _add_inline_runs(p, text, size_pt=10.5)

        # Blockquote / comentario TODO — omitir silenciosamente
        elif line.strip().startswith(">") or line.strip().startswith("<!--"):
            pass

        # Linea vacia
        elif line.strip() == "":
            pass

        # Texto normal
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _add_inline_runs(p, line.strip(), size_pt=10.5)

        i += 1


def _add_inline_runs(paragraph, text, size_pt=10.5):
    """
    Parsea inline markdown en un parrafo:
    - **texto** → run bold
    - `codigo` → run plain (sin backticks)
    - [texto](url) → muestra solo el texto del enlace
    - §        → 'seccion'
    Todo lo demas → run normal.
    """
    text = text.replace("§", "sección")
    # Patron: [text](url), **bold**, *italic*, `code`
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))", text)
    for token in tokens:
        if not token:
            continue
        # Hipervínculo [text](url)
        m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
        if m:
            E.add_hyperlink(paragraph, m.group(1), m.group(2))
            continue
        bold = False
        italic = False
        if token.startswith("**") and token.endswith("**"):
            token = token[2:-2]
            bold = True
        elif token.startswith("*") and token.endswith("*"):
            token = token[1:-1]
            italic = True
        elif token.startswith("`") and token.endswith("`"):
            token = token[1:-1]   # strip backticks — show as plain
        run = paragraph.add_run(token)
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = E.COLOR_TEXT
        E._safe_font(run)


def _clean_cell(text):
    """Limpia markdown inline de celdas de tabla."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # [text](url) → text
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.replace("§", "sección").strip()


def _render_table(doc, pipe_lines):
    """Renderiza tabla desde lineas pipe de Markdown."""
    rows_data = []
    for line in pipe_lines:
        cells = [_clean_cell(c) for c in line.strip().strip("|").split("|")]
        rows_data.append(cells)

    if not rows_data:
        return

    headers = rows_data[0]
    data    = rows_data[1:] if len(rows_data) > 1 else []
    E.add_table(doc, headers=headers, rows=data)
    doc.add_paragraph()


# ── Header y footer ───────────────────────────────────────────────────────────
def add_header_footer(doc, meta):
    footer_text = f"{meta.get('subtitulo', '')} — Plan de Pruebas v{meta.get('version', '1.0')}"
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Texto izquierdo
        run_left = p.add_run(footer_text + "    ")
        run_left.font.size = Pt(7.5)
        run_left.font.color.rgb = E.COLOR_MUTED
        E._safe_font(run_left)

        # "Pag. X de Y" como texto + campos OOXML en runs separados
        _append_page_of_total(p)


def _append_page_of_total(paragraph):
    """Agrega 'Pag. X de Y' al final del parrafo via OOXML."""
    def _field_run(p, instr_text):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "15")   # 7.5pt * 2
        rPr.append(sz)
        r.append(rPr)

        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        r.append(begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instr_text
        r.append(instr)

        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        r.append(sep)

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r.append(end)

        paragraph._p.append(r)

    def _text_run(p, text):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "15")
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        paragraph._p.append(r)

    _text_run(paragraph, "Pag. ")
    _field_run(paragraph, " PAGE ")
    _text_run(paragraph, " de ")
    _field_run(paragraph, " NUMPAGES ")


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    # Leer metadata
    with open(CONTENIDO / "00-portada.yaml", encoding="utf-8") as f:
        meta = yaml.safe_load(f)

    doc = Document()

    # Margenes globales
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Estructura del documento ──────────────────────────────────────────────
    build_portada(doc, meta)
    build_toc(doc)
    build_historial(doc, meta)
    build_info_proyecto(doc, meta)

    # ── Secciones desde Markdown ──────────────────────────────────────────────
    for filename in SECCIONES:
        path = CONTENIDO / filename
        if not path.exists():
            print(f"WARN: {filename} no encontrado, omitiendo.")
            continue
        text = path.read_text(encoding="utf-8")
        parse_markdown_section(doc, text)
        doc.add_paragraph()

    # ── Header y footer ───────────────────────────────────────────────────────
    add_header_footer(doc, meta)

    # ── Forzar actualizacion de campos (TOC) al abrir en Word ────────────────
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    # ── Guardar ───────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    main()
