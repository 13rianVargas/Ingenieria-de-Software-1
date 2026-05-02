"""
Paleta charcoal/slate y helpers de estilo para python-docx.
Coherente con la paleta del Taller-6 (arquitectura.pdf).
"""

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Paleta ────────────────────────────────────────────────────────────────────
COLOR_PRIMARY   = RGBColor(0x2B, 0x2D, 0x42)  # deep slate  → titulos, encabezados tabla
COLOR_SECONDARY = RGBColor(0x55, 0x5B, 0x6E)  # medium slate → h2, subtitulos
COLOR_LIGHT     = RGBColor(0xE8, 0xEA, 0xEF)  # pale slate  → filas alternas tabla
COLOR_BORDER    = RGBColor(0xA8, 0xAC, 0xBA)  # medium gray → bordes tabla
COLOR_MUTED     = RGBColor(0x6C, 0x70, 0x80)  # gray        → pie de pagina, leyendas
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TEXT      = RGBColor(0x21, 0x25, 0x29)  # casi negro  → cuerpo


# ── Fuente base ───────────────────────────────────────────────────────────────
FONT_FAMILY = "Helvetica Neue"
FONT_FAMILY_FALLBACK = "Calibri"  # fallback si HN no esta instalada


def _safe_font(run):
    """Aplica fuente con fallback."""
    run.font.name = FONT_FAMILY
    run._element.rPr.rFonts.set(qn("w:cs"), FONT_FAMILY_FALLBACK)


# ── Helpers de parrafo ────────────────────────────────────────────────────────
def set_paragraph_style(paragraph, font_size_pt, bold=False, color=COLOR_TEXT,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before_pt=0,
                        space_after_pt=6):
    """Aplica estilo basico a un parrafo completo."""
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after  = Pt(space_after_pt)
    for run in paragraph.runs:
        run.font.size  = Pt(font_size_pt)
        run.font.bold  = bold
        run.font.color.rgb = color
        _safe_font(run)


def add_heading(doc, text, level=1):
    """Agrega heading con estilo charcoal. Usa Heading 1/2/3 nativo para que el TOC lo detecte."""
    p = doc.add_paragraph(style=f'Heading {level}')
    run = p.add_run(text)
    sizes   = {1: 16, 2: 13, 3: 11}
    colors  = {1: COLOR_PRIMARY, 2: COLOR_SECONDARY, 3: COLOR_SECONDARY}
    run.font.size  = Pt(sizes.get(level, 11))
    run.font.bold  = True
    run.font.color.rgb = colors.get(level, COLOR_PRIMARY)
    _safe_font(run)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    if level == 1:
        # linea inferior decorativa
        _add_bottom_border(p, color="2B2D42", size=8)
    return p


def add_body(doc, text):
    """Agrega parrafo de cuerpo estandar."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_TEXT
        _safe_font(run)
    p.paragraph_format.space_after = Pt(6)
    return p


# ── Helpers de tabla ─────────────────────────────────────────────────────────
def add_table(doc, headers, rows, col_widths_cm=None):
    """
    Agrega tabla con estilo charcoal.
    headers: list[str]
    rows: list[list[str]]
    col_widths_cm: list[float] | None (distribucion uniforme si None)
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Encabezado
    hdr_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = text
        _style_header_cell(cell)

    # Filas de datos
    for r_idx, row_data in enumerate(rows):
        data_row = table.rows[r_idx + 1]
        bg = COLOR_LIGHT if r_idx % 2 == 0 else COLOR_WHITE
        for c_idx, text in enumerate(row_data):
            cell = data_row.cells[c_idx]
            cell.text = text
            _style_data_cell(cell, bg)

    # Anchos de columna
    if col_widths_cm:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths_cm):
                    cell.width = Cm(col_widths_cm[i])

    return table


def _style_header_cell(cell):
    _set_cell_bg(cell, "2B2D42")
    p = cell.paragraphs[0]
    for run in p.runs:
        run.font.bold  = True
        run.font.size  = Pt(9)
        run.font.color.rgb = COLOR_WHITE
        _safe_font(run)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


def _style_data_cell(cell, bg_color=COLOR_WHITE):
    _set_cell_bg(cell, "{:02X}{:02X}{:02X}".format(*bg_color))
    p = cell.paragraphs[0]
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_TEXT
        _safe_font(run)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)


# ── OOXML helpers ─────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color):
    """Establece color de fondo de celda via OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_bottom_border(paragraph, color="2B2D42", size=8):
    """Agrega borde inferior a un parrafo via OOXML."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hyperlink(paragraph, text, url):
    """Agrega un run con hipervínculo real al párrafo (OOXML)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


def add_toc_field(doc):
    """
    Inserta campo TOC via OOXML.
    Word/LibreOffice lo regenera al abrir el archivo (F9 o al abrir).
    """
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    run._r.append(fldChar_end)
    return p
